"""Extract raw text from resume files."""
import logging
from pathlib import Path

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class UnreadableFileError(Exception):
    """Raised when a resume file cannot be read or parsed."""


def read_pdf(file_path: Path) -> str:
    try:
        reader = PdfReader(file_path)
    except Exception as exc:
        raise UnreadableFileError(f"Could not open PDF {file_path.name}: {exc}") from exc

    text_parts = []
    for page_num, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text()
        except Exception as exc:
            logger.warning("Failed to extract page %d of %s: %s", page_num, file_path.name, exc)
            continue
        if page_text:
            text_parts.append(page_text)

    text = "\n".join(text_parts)
    if not text.strip():
        raise UnreadableFileError(f"No extractable text in {file_path.name} (likely scanned/image-based).")
    return text


def read_docx(file_path: Path) -> str:
    try:
        document = Document(file_path)
    except Exception as exc:
        raise UnreadableFileError(f"Could not open DOCX {file_path.name}: {exc}") from exc

    text_parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    text = "\n".join(text_parts)
    if not text.strip():
        raise UnreadableFileError(f"No extractable text in {file_path.name}.")
    return text


def read_resume(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(file_path)
    elif suffix == ".docx":
        return read_docx(file_path)
    else:
        raise UnreadableFileError(f"Unsupported file type: {suffix}")
